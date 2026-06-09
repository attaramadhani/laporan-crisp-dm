"""
laporan_crisp_dm_v5.py
======================
Mode: SILENT (verbose=0, warnings suppressed)
Tujuan: Pipeline ML lengkap dengan perbaikan metodologi dari v2:
        - SMOTE dimasukkan ke dalam imblearn.Pipeline agar tidak bocor ke CV
        - N_ITER ditingkatkan (50 RF, 60 XGB) untuk pencarian yang lebih luas
        - max_depth=None dihapus dari search space untuk mencegah overfitting
        - Baseline (tanpa tuning) + Tuned (dengan pipeline CV yang benar)
        - Laporan Word format CRISP-DM lengkap

Hardware target: MSI Cyborg 15 A13V (i5-13420H, 12 threads)
Dataset        : final_dataset_kspr_attala.csv
"""

import warnings
warnings.filterwarnings('ignore')

import os, sys, time, joblib
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO

from sklearn.model_selection import train_test_split, RandomizedSearchCV, StratifiedKFold
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import (classification_report, roc_auc_score,
                             confusion_matrix, f1_score, accuracy_score,
                             roc_curve, auc)
from sklearn.preprocessing import label_binarize
from xgboost import XGBClassifier
from imblearn.over_sampling import SMOTE
from imblearn.pipeline import Pipeline as ImbPipeline
import shap

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx tidak terinstall. Jalankan: pip install python-docx")
    sys.exit(1)

# =========================================================== #
#  KONSTANTA KONFIGURASI                                      #
# =========================================================== #
N_JOBS       = 10     # 10 dari 12 thread
RANDOM_STATE = 42
N_ITER_RF    = 50     # ditingkatkan dari 20 → 50 (lebih luas)
N_ITER_XGB   = 60     # ditingkatkan dari 25 → 60 (lebih luas)
OUTPUT_FILE  = 'output_eksperimen/Laporan_CRISP_DM_Komputasi_v5.docx'
CACHE_FILE   = 'output_eksperimen/v5_cache.pkl'
DATA_FILE    = 'final_dataset_kspr_attala.csv'

class_labels = {0: 'Low', 1: 'Moderate', 2: 'High'}

def log(msg):
    print(f"[{time.strftime('%H:%M:%S')}] {msg}", flush=True)

# =========================================================== #
#  1. LOAD DATA                                               #
# =========================================================== #
log("1/6 Loading data...")
try:
    df = pd.read_csv(DATA_FILE)
except FileNotFoundError:
    log(f"ERROR: File '{DATA_FILE}' not found.")
    sys.exit(1)
except Exception as e:
    log(f"ERROR reading file: {e}")
    sys.exit(1)

n_rows_raw      = len(df)
n_cols_raw      = len(df.columns)
target_dist_raw = df['label_risiko'].value_counts().sort_index().to_dict()
missing_total   = int(df.isnull().sum().sum())

# =========================================================== #
#  2. PREPROCESSING                                           #
# =========================================================== #
log("2/6 Feature Selection (Removing Leakage/ID features)...")
drop_cols = ['label_risiko']

id_cols = ['id_anggota_rt', 'id_rumah_tangga', 'id_provinsi', 'id_kabupaten']
for c in id_cols:
    if c in df.columns:
        drop_cols.append(c)

if 'metode_persalinan_sesar' in df.columns: drop_cols.append('metode_persalinan_sesar')
if 'operasi_caesar' in df.columns:          drop_cols.append('operasi_caesar')

X = df.drop(drop_cols, axis=1)
y = df['label_risiko']
n_features = X.shape[1]

# =========================================================== #
#  3. TRAIN-TEST SPLIT                                        #
#     SMOTE tidak dilakukan di sini — hanya split saja        #
#     SMOTE akan masuk ke dalam pipeline CV                   #
# =========================================================== #
log("3/6 Train-Test Split (80/20 Stratified)...")
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.20, random_state=RANDOM_STATE, stratify=y
)
dist_train_original = y_train.value_counts().sort_index().to_dict()

# Hitung distribusi SMOTE untuk dokumentasi (simulasi di luar CV)
_smote_doc = SMOTE(random_state=RANDOM_STATE)
_, y_train_s_doc = _smote_doc.fit_resample(X_train, y_train)
dist_after_smote = pd.Series(y_train_s_doc).value_counts().sort_index().to_dict()
del _smote_doc, y_train_s_doc

os.makedirs('output_eksperimen', exist_ok=True)

# =========================================================== #
#  CACHING: Load hasil jika sudah pernah dihitung             #
# =========================================================== #
if os.path.exists(CACHE_FILE):
    log(f"[CACHE] Cache '{CACHE_FILE}' found -> skipping training & tuning, loading directly.")
    cache = joblib.load(CACHE_FILE)
    # Baseline
    acc_rf_b      = cache['acc_rf_b']
    f1_rf_b       = cache['f1_rf_b']
    roc_rf_b      = cache['roc_rf_b']
    report_rf_b   = cache['report_rf_b']
    cm_rf_b       = cache['cm_rf_b']
    y_pred_rf_b   = cache['y_pred_rf_b']
    y_proba_rf_b  = cache['y_proba_rf_b']
    time_rf_base  = cache['time_rf_base']

    acc_xgb_b     = cache['acc_xgb_b']
    f1_xgb_b      = cache['f1_xgb_b']
    roc_xgb_b     = cache['roc_xgb_b']
    report_xgb_b  = cache['report_xgb_b']
    cm_xgb_b      = cache['cm_xgb_b']
    y_pred_xgb_b  = cache['y_pred_xgb_b']
    y_proba_xgb_b = cache['y_proba_xgb_b']
    time_xgb_base = cache['time_xgb_base']

    # Tuned
    acc_rf_t      = cache['acc_rf_t']
    f1_rf_t       = cache['f1_rf_t']
    roc_rf_t      = cache['roc_rf_t']
    report_rf_t   = cache['report_rf_t']
    cm_rf_t       = cache['cm_rf_t']
    y_pred_rf_t   = cache['y_pred_rf_t']
    y_proba_rf_t  = cache['y_proba_rf_t']
    rf_tuned      = cache['rf_tuned']
    rf_best       = cache['rf_best']
    time_rf_tune  = cache['time_rf_tune']

    acc_xgb_t     = cache['acc_xgb_t']
    f1_xgb_t      = cache['f1_xgb_t']
    roc_xgb_t     = cache['roc_xgb_t']
    report_xgb_t  = cache['report_xgb_t']
    cm_xgb_t      = cache['cm_xgb_t']
    y_pred_xgb_t  = cache['y_pred_xgb_t']
    y_proba_xgb_t = cache['y_proba_xgb_t']
    xgb_tuned     = cache['xgb_tuned']
    xgb_best      = cache['xgb_best']
    time_xgb_tune = cache['time_xgb_tune']

    log(f"  RF  Baseline -> Acc={acc_rf_b:.4f}  F1={f1_rf_b:.4f}  ROC={roc_rf_b:.4f}  (from cache)")
    log(f"  XGB Baseline -> Acc={acc_xgb_b:.4f}  F1={f1_xgb_b:.4f}  ROC={roc_xgb_b:.4f}  (from cache)")
    log(f"  RF  Tuned    -> CV-best=cached  Test ROC={roc_rf_t:.4f}  (from cache)")
    log(f"  XGB Tuned    -> CV-best=cached  Test ROC={roc_xgb_t:.4f}  (from cache)")

else:
    # =========================================================== #
    #  4. BASELINE MODELS (SMOTE di luar, tanpa tuning)           #
    # =========================================================== #
    log("4/6 Defining ImbPipeline (SMOTE + Model)...")

    smote_base = SMOTE(random_state=RANDOM_STATE)
    X_train_s, y_train_s = smote_base.fit_resample(X_train, y_train)

    # --- Random Forest Baseline ---
    t0 = time.time()
    rf_base = RandomForestClassifier(random_state=RANDOM_STATE, n_jobs=N_JOBS)
    rf_base.fit(X_train_s, y_train_s)
    time_rf_base = time.time() - t0

    y_pred_rf_b  = rf_base.predict(X_test)
    y_proba_rf_b = rf_base.predict_proba(X_test)
    acc_rf_b     = accuracy_score(y_test, y_pred_rf_b)
    f1_rf_b      = f1_score(y_test, y_pred_rf_b, average='macro')
    roc_rf_b     = roc_auc_score(y_test, y_proba_rf_b, multi_class='ovr')
    report_rf_b  = classification_report(y_test, y_pred_rf_b, output_dict=True)
    cm_rf_b      = confusion_matrix(y_test, y_pred_rf_b)

    # --- XGBoost Baseline ---
    t0 = time.time()
    xgb_base = XGBClassifier(
        random_state=RANDOM_STATE, objective='multi:softprob',
        eval_metric='mlogloss', n_jobs=N_JOBS, tree_method='hist', verbosity=0
    )
    xgb_base.fit(X_train_s, y_train_s)
    time_xgb_base = time.time() - t0

    y_pred_xgb_b  = xgb_base.predict(X_test)
    y_proba_xgb_b = xgb_base.predict_proba(X_test)
    acc_xgb_b     = accuracy_score(y_test, y_pred_xgb_b)
    f1_xgb_b      = f1_score(y_test, y_pred_xgb_b, average='macro')
    roc_xgb_b     = roc_auc_score(y_test, y_proba_xgb_b, multi_class='ovr')
    report_xgb_b  = classification_report(y_test, y_pred_xgb_b, output_dict=True)
    cm_xgb_b      = confusion_matrix(y_test, y_pred_xgb_b)

    log(f"  RF  Baseline -> Acc={acc_rf_b:.4f}  F1={f1_rf_b:.4f}  ROC={roc_rf_b:.4f}  ({time_rf_base:.2f}s)")
    log(f"  XGB Baseline -> Acc={acc_xgb_b:.4f}  F1={f1_xgb_b:.4f}  ROC={roc_xgb_b:.4f}  ({time_xgb_base:.2f}s)")

    # =========================================================== #
    #  5. HYPERPARAMETER TUNING DENGAN IMBPIPELINE                #
    # =========================================================== #
    log("5/6 Hyperparameter Tuning (ImbPipeline + RandomizedSearchCV, mode silent)...")

    cv_strat = StratifiedKFold(n_splits=5, shuffle=True, random_state=RANDOM_STATE)

    # -- RF Pipeline + Param Space --
    rf_pipe = ImbPipeline([
        ('smote', SMOTE(random_state=RANDOM_STATE)),
        ('clf',   RandomForestClassifier(random_state=RANDOM_STATE, n_jobs=1))
    ])
    rf_param_dist = {
        'clf__n_estimators':      [100, 150, 200, 250, 300],
        'clf__max_depth':         [10, 15, 20, 25],
        'clf__min_samples_split': [2, 5, 10],
        'clf__min_samples_leaf':  [1, 2, 4],
        'clf__max_features':      ['sqrt', 'log2'],
        'clf__bootstrap':         [True],
    }
    log(f"  Tuning RF (ImbPipeline): {N_ITER_RF} iter x 5 fold = {N_ITER_RF*5} fits ...")
    t0 = time.time()
    rf_rs = RandomizedSearchCV(
        rf_pipe, param_distributions=rf_param_dist,
        n_iter=N_ITER_RF, cv=cv_strat, scoring='f1_macro',
        n_jobs=N_JOBS, pre_dispatch='2*n_jobs',
        random_state=RANDOM_STATE, verbose=0, refit=True
    )
    rf_rs.fit(X_train, y_train)
    time_rf_tune = time.time() - t0

    rf_tuned     = rf_rs.best_estimator_
    y_pred_rf_t  = rf_tuned.predict(X_test)
    y_proba_rf_t = rf_tuned.predict_proba(X_test)
    acc_rf_t     = accuracy_score(y_test, y_pred_rf_t)
    f1_rf_t      = f1_score(y_test, y_pred_rf_t, average='macro')
    roc_rf_t     = roc_auc_score(y_test, y_proba_rf_t, multi_class='ovr')
    report_rf_t  = classification_report(y_test, y_pred_rf_t, output_dict=True)
    cm_rf_t      = confusion_matrix(y_test, y_pred_rf_t)
    rf_best      = {k.replace('clf__', ''): v for k, v in rf_rs.best_params_.items()}
    log(f"  RF  Tuned  -> CV-best={rf_rs.best_score_:.4f}  Test ROC={roc_rf_t:.4f}  ({time_rf_tune/60:.1f} mnt)")

    # -- XGB Pipeline + Param Space --
    xgb_pipe = ImbPipeline([
        ('smote', SMOTE(random_state=RANDOM_STATE)),
        ('clf',   XGBClassifier(
            random_state=RANDOM_STATE, objective='multi:softprob',
            eval_metric='mlogloss', tree_method='hist', verbosity=0, n_jobs=1
        ))
    ])
    xgb_param_dist = {
        'clf__n_estimators':     [100, 150, 200, 250, 300],
        'clf__max_depth':        [3, 4, 5, 6, 7],
        'clf__learning_rate':    [0.01, 0.05, 0.1, 0.2],
        'clf__subsample':        [0.7, 0.8, 0.9, 1.0],
        'clf__colsample_bytree': [0.7, 0.8, 0.9, 1.0],
        'clf__gamma':            [0, 0.1, 0.2],
        'clf__min_child_weight': [1, 3, 5],
        'clf__reg_alpha':        [0, 0.1, 0.5],
        'clf__reg_lambda':       [0.5, 1.0, 1.5],
    }
    log(f"  Tuning XGB (ImbPipeline): {N_ITER_XGB} iter x 5 fold = {N_ITER_XGB*5} fits ...")
    t0 = time.time()
    xgb_rs = RandomizedSearchCV(
        xgb_pipe, param_distributions=xgb_param_dist,
        n_iter=N_ITER_XGB, cv=cv_strat, scoring='f1_macro',
        n_jobs=N_JOBS, pre_dispatch='2*n_jobs',
        random_state=RANDOM_STATE, verbose=0, refit=True
    )
    xgb_rs.fit(X_train, y_train)
    time_xgb_tune = time.time() - t0

    xgb_tuned     = xgb_rs.best_estimator_
    y_pred_xgb_t  = xgb_tuned.predict(X_test)
    y_proba_xgb_t = xgb_tuned.predict_proba(X_test)
    acc_xgb_t     = accuracy_score(y_test, y_pred_xgb_t)
    f1_xgb_t      = f1_score(y_test, y_pred_xgb_t, average='macro')
    roc_xgb_t     = roc_auc_score(y_test, y_proba_xgb_t, multi_class='ovr')
    report_xgb_t  = classification_report(y_test, y_pred_xgb_t, output_dict=True)
    cm_xgb_t      = confusion_matrix(y_test, y_pred_xgb_t)
    xgb_best      = {k.replace('clf__', ''): v for k, v in xgb_rs.best_params_.items()}
    log(f"  XGB Tuned  -> CV-best={xgb_rs.best_score_:.4f}  Test ROC={roc_xgb_t:.4f}  ({time_xgb_tune/60:.1f} mnt)")

    # =========================================================== #
    #  SIMPAN CACHE                                               #
    # =========================================================== #
    log(f"[CACHE] Menyimpan hasil ke '{CACHE_FILE}'...")
    joblib.dump({
        # Baseline RF
        'acc_rf_b': acc_rf_b, 'f1_rf_b': f1_rf_b, 'roc_rf_b': roc_rf_b,
        'report_rf_b': report_rf_b, 'cm_rf_b': cm_rf_b,
        'y_pred_rf_b': y_pred_rf_b, 'y_proba_rf_b': y_proba_rf_b,
        'time_rf_base': time_rf_base,
        # Baseline XGB
        'acc_xgb_b': acc_xgb_b, 'f1_xgb_b': f1_xgb_b, 'roc_xgb_b': roc_xgb_b,
        'report_xgb_b': report_xgb_b, 'cm_xgb_b': cm_xgb_b,
        'y_pred_xgb_b': y_pred_xgb_b, 'y_proba_xgb_b': y_proba_xgb_b,
        'time_xgb_base': time_xgb_base,
        # Tuned RF
        'acc_rf_t': acc_rf_t, 'f1_rf_t': f1_rf_t, 'roc_rf_t': roc_rf_t,
        'report_rf_t': report_rf_t, 'cm_rf_t': cm_rf_t,
        'y_pred_rf_t': y_pred_rf_t, 'y_proba_rf_t': y_proba_rf_t,
        'rf_tuned': rf_tuned, 'rf_best': rf_best, 'time_rf_tune': time_rf_tune,
        # Tuned XGB
        'acc_xgb_t': acc_xgb_t, 'f1_xgb_t': f1_xgb_t, 'roc_xgb_t': roc_xgb_t,
        'report_xgb_t': report_xgb_t, 'cm_xgb_t': cm_xgb_t,
        'y_pred_xgb_t': y_pred_xgb_t, 'y_proba_xgb_t': y_proba_xgb_t,
        'xgb_tuned': xgb_tuned, 'xgb_best': xgb_best, 'time_xgb_tune': time_xgb_tune,
    }, CACHE_FILE)
    log(f"[CACHE] Tersimpan. Run ulang berikutnya akan langsung pakai cache ini.")

# =========================================================== #
#  PLOTS                                                      #
# =========================================================== #
def fig_to_bytes(fig):
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    return buf

# ── Confusion Matrices (4 model) ──
fig_cm, axes = plt.subplots(2, 2, figsize=(14, 10))
plot_data = [
    (cm_rf_b,  'Random Forest – Baseline'),
    (cm_xgb_b, 'XGBoost – Baseline'),
    (cm_rf_t,  'Random Forest – Tuned (Pipeline)'),
    (cm_xgb_t, 'XGBoost – Tuned (Pipeline)'),
]
for ax, (cm, ttl) in zip(axes.flatten(), plot_data):
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', ax=ax,
                xticklabels=['Low','Moderate','High'],
                yticklabels=['Low','Moderate','High'])
    ax.set_title(ttl, fontsize=11, fontweight='bold')
    ax.set_xlabel('Prediksi'); ax.set_ylabel('Aktual')
plt.suptitle('Confusion Matrix – Baseline vs Tuned (ImbPipeline)', fontsize=13, fontweight='bold')
plt.tight_layout()
cm_buf = fig_to_bytes(fig_cm)
plt.close()

# ── Feature Importance (dari classifier dalam pipeline) ──
fig_fi, axes = plt.subplots(1, 2, figsize=(16, 7))
for ax, pipe_model, color, ttl in zip(
    axes,
    [rf_tuned, xgb_tuned],
    ['steelblue', 'forestgreen'],
    ['Random Forest (Tuned)', 'XGBoost (Tuned)'],
):
    clf_step = pipe_model.named_steps['clf']
    imp = clf_step.feature_importances_
    idx = np.argsort(imp)[-15:]
    ax.barh(range(len(idx)), imp[idx], color=color, align='center')
    ax.set_yticks(range(len(idx)))
    ax.set_yticklabels([X.columns[i] for i in idx], fontsize=9)
    ax.set_xlabel('Importance Score')
    ax.set_title(f'Top 15 Feature Importance\n{ttl}', fontweight='bold')
plt.tight_layout()
fi_buf = fig_to_bytes(fig_fi)
plt.close()

# ── ROC-AUC Comparison Bar Chart ──
fig_roc, ax = plt.subplots(figsize=(11, 5))
models_lbl = ['RF\nBaseline', 'XGB\nBaseline', 'RF\nTuned', 'XGB\nTuned']
roc_scores = [roc_rf_b, roc_xgb_b, roc_rf_t, roc_xgb_t]
colors_roc = ['#5885AF', '#3E7A5E', '#1C4E80', '#0C5928']
bars = ax.bar(models_lbl, roc_scores, color=colors_roc, width=0.5)
ax.set_ylim(max(0.80, min(roc_scores) - 0.03), 1.0)
ax.set_ylabel('ROC-AUC Score (OvR)')
ax.set_title('Perbandingan ROC-AUC Score – Baseline vs Tuned (ImbPipeline)', fontweight='bold')
for bar, score in zip(bars, roc_scores):
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.001,
            f'{score:.4f}', ha='center', va='bottom', fontweight='bold')
plt.tight_layout()
roc_buf = fig_to_bytes(fig_roc)
plt.close()

# ── ROC Curve (OvR per kelas) – semua 4 model ──
log("Plot kurva ROC (OvR per kelas, 4 model)...")

classes_list  = [0, 1, 2]
class_colors  = ['#e84393', '#f59e0b', '#10b981']
class_labels_roc = ['Low (0)', 'Moderate (1)', 'High (2)']

y_test_bin = label_binarize(y_test, classes=classes_list)

_roc_models = [
    ('RF Baseline',          y_proba_rf_b,  '#5885AF'),
    ('XGBoost Baseline',     y_proba_xgb_b, '#3E7A5E'),
    ('RF Tuned (Pipeline)',  y_proba_rf_t,  '#1C4E80'),
    ('XGB Tuned (Pipeline)', y_proba_xgb_t, '#0C5928'),
]

# 4 sub-plot: satu per model, setiap sub-plot ada 3 kurva kelas
fig_roc_curve, axes_roc = plt.subplots(2, 2, figsize=(14, 11))
for ax, (model_name, y_proba, base_color) in zip(axes_roc.flatten(), _roc_models):
    for i, (cls, color, lbl) in enumerate(zip(classes_list, class_colors, class_labels_roc)):
        fpr, tpr, _ = roc_curve(y_test_bin[:, i], y_proba[:, i])
        roc_auc_val = auc(fpr, tpr)
        ax.plot(fpr, tpr, color=color, lw=2,
                label=f'{lbl} (AUC = {roc_auc_val:.3f})')
    ax.plot([0, 1], [0, 1], 'k--', lw=1, alpha=0.6, label='Random Classifier')
    ax.set_xlim([0.0, 1.0])
    ax.set_ylim([0.0, 1.05])
    ax.set_xlabel('False Positive Rate', fontsize=10)
    ax.set_ylabel('True Positive Rate', fontsize=10)
    ax.set_title(f'ROC Curve – {model_name}', fontsize=11, fontweight='bold')
    ax.legend(loc='lower right', fontsize=9)
    ax.grid(True, alpha=0.3)
plt.suptitle('ROC Curve One-vs-Rest (OvR) – Baseline vs Tuned (ImbPipeline)',
             fontsize=13, fontweight='bold')
plt.tight_layout()
roc_curve_buf = fig_to_bytes(fig_roc_curve)
plt.close()

# ── SHAP (dari classifier tuned) ──
log("Kalkulasi SHAP values (subset 300 untuk efisiensi)...")
X_test_sample = (X_test.sample(min(300, len(X_test)), random_state=RANDOM_STATE)
                 if len(X_test) > 300 else X_test)

rf_clf_step  = rf_tuned.named_steps['clf']
xgb_clf_step = xgb_tuned.named_steps['clf']

# SHAP RF
explainer_rf   = shap.TreeExplainer(rf_clf_step)
shap_values_rf = explainer_rf.shap_values(X_test_sample)

plt.figure()
shap.summary_plot(shap_values_rf, X_test_sample, plot_type="bar", show=False)
plt.title("SHAP Global Feature Importance (Random Forest Tuned)", pad=20)
plt.tight_layout()
shap_rf_buf = fig_to_bytes(plt.gcf())
plt.close()

if isinstance(shap_values_rf, list) and len(shap_values_rf) > 2:
    plt.figure()
    shap.summary_plot(shap_values_rf[2], X_test_sample, show=False)
    plt.title("Random Forest SHAP – Risiko TINGGI", pad=20)
    plt.tight_layout()
    shap_rf_tinggi_buf = fig_to_bytes(plt.gcf())
    plt.close()
elif hasattr(shap_values_rf, 'shape') and len(shap_values_rf.shape) == 3:
    plt.figure()
    shap.summary_plot(shap_values_rf[:, :, 2], X_test_sample, show=False)
    plt.title("Random Forest SHAP – Risiko TINGGI", pad=20)
    plt.tight_layout()
    shap_rf_tinggi_buf = fig_to_bytes(plt.gcf())
    plt.close()
else:
    shap_rf_tinggi_buf = None

# SHAP XGB
explainer_xgb   = shap.TreeExplainer(xgb_clf_step)
shap_values_xgb = explainer_xgb.shap_values(X_test_sample)

plt.figure()
shap.summary_plot(shap_values_xgb, X_test_sample, plot_type="bar", show=False)
plt.title("SHAP Global Feature Importance (XGBoost Tuned)", pad=20)
plt.tight_layout()
shap_xgb_buf = fig_to_bytes(plt.gcf())
plt.close()

if isinstance(shap_values_xgb, list) and len(shap_values_xgb) > 2:
    plt.figure()
    shap.summary_plot(shap_values_xgb[2], X_test_sample, show=False)
    plt.title("XGBoost SHAP – Risiko TINGGI", pad=20)
    plt.tight_layout()
    shap_xgb_tinggi_buf = fig_to_bytes(plt.gcf())
    plt.close()
elif hasattr(shap_values_xgb, 'shape') and len(shap_values_xgb.shape) == 3:
    plt.figure()
    shap.summary_plot(shap_values_xgb[:, :, 2], X_test_sample, show=False)
    plt.title("XGBoost SHAP – Risiko TINGGI", pad=20)
    plt.tight_layout()
    shap_xgb_tinggi_buf = fig_to_bytes(plt.gcf())
    plt.close()
else:
    shap_xgb_tinggi_buf = None

# =========================================================== #
#  6. BUAT DOKUMEN WORD                                       #
# =========================================================== #
log("6/6 Membuat dokumen Word...")
os.makedirs('output_eksperimen', exist_ok=True)

doc = Document()
sec = doc.sections[0]
sec.page_width    = Cm(21);  sec.page_height   = Cm(29.7)
sec.left_margin   = Cm(3);   sec.right_margin  = Cm(2.5)
sec.top_margin    = Cm(2.5); sec.bottom_margin = Cm(2.5)

def h(lvl, txt):
    p = doc.add_heading(txt, level=lvl)
    p.paragraph_format.space_before = Pt(10 if lvl == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)
    return p

def para(txt, bold=False, italic=False, align=None):
    p = doc.add_paragraph(txt)
    if bold or italic:
        for run in p.runs:
            run.bold   = bold
            run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(txt):
    doc.add_paragraph(txt, style='List Bullet')

def add_table(headers, rows_data):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    hd_row = tbl.rows[0].cells
    for i, h_txt in enumerate(headers):
        hd_row[i].text = h_txt
        for run in hd_row[i].paragraphs[0].runs:
            run.bold = True
    for row_vals in rows_data:
        row = tbl.add_row().cells
        for i, val in enumerate(row_vals):
            row[i].text = str(val)
    doc.add_paragraph()
    return tbl

def add_monospaced_block(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(10)
    return p

# ═══════════════════════════════════════════════════════════ #
#  TITLE PAGE                                                 #
# ═══════════════════════════════════════════════════════════ #
h(0, '')
p_title = doc.add_heading('Research Report (Version 5)', 0)
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p_sub = doc.add_heading(
    'Pregnancy Risk Classification using\n'
    'Random Forest & XGBoost\n'
    'SMOTE within ImbPipeline + Hyperparameter Tuning (Correct Methodology)\n'
    '(Final Dataset: KSPR Attala)', 1)
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

for txt in [
    'Framework: CRISP-DM (Cross-Industry Standard Process for Data Mining)',
    'Hardware: MSI Cyborg 15 A13V (Intel Core i5-13420H)',
    f'Experiment Date: {time.strftime("%d %B %Y")}',
]:
    para(txt, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════ #
#  EXECUTIVE SUMMARY                                          #
# ═══════════════════════════════════════════════════════════ #
h(1, 'Executive Summary')

best_scores = {
    'RF Baseline':  roc_rf_b,
    'XGB Baseline': roc_xgb_b,
    'RF Tuned':     roc_rf_t,
    'XGB Tuned':    roc_xgb_t,
}
best_name = max(best_scores, key=best_scores.get)
best_roc  = best_scores[best_name]
best_f1   = {'RF Baseline': f1_rf_b, 'XGB Baseline': f1_xgb_b,
             'RF Tuned': f1_rf_t, 'XGB Tuned': f1_xgb_t}[best_name]

para(
    f'This Version 5 research implements critical methodological improvements over Version 2: '
    f'SMOTE is integrated into imblearn.Pipeline so that resampling only occurs '
    f'on the training fold — avoiding data leakage to the validation fold '
    f'during cross-validation. N_ITER is increased to {N_ITER_RF} (RF) and '
    f'{N_ITER_XGB} (XGB) for a broader hyperparameter search. '
    f'The final dataset final_dataset_kspr_attala.csv with {n_rows_raw:,} observations was used. '
    f'The best model is {best_name} with an ROC-AUC = {best_roc:.4f}.'
)

add_table(
    ['Metric', 'RF Baseline', 'XGB Baseline', 'RF Tuned', 'XGB Tuned'],
    [
        ('Accuracy',  f'{acc_rf_b:.4f}',  f'{acc_xgb_b:.4f}',  f'{acc_rf_t:.4f}',  f'{acc_xgb_t:.4f}'),
        ('F1-Macro',  f'{f1_rf_b:.4f}',   f'{f1_xgb_b:.4f}',   f'{f1_rf_t:.4f}',   f'{f1_xgb_t:.4f}'),
        ('ROC-AUC',   f'{roc_rf_b:.4f}',  f'{roc_xgb_b:.4f}',  f'{roc_rf_t:.4f}',  f'{roc_xgb_t:.4f}'),
        ('Fit Time', f'{time_rf_base:.2f}s', f'{time_xgb_base:.2f}s',
                      f'{time_rf_tune/60:.1f}m', f'{time_xgb_tune/60:.1f}m'),
    ]
)
doc.add_picture(roc_buf, width=Inches(5.5))
doc.add_page_break()

# ═══════════════════════════════════════════════════════════ #
#  CHAPTER 1: BUSINESS UNDERSTANDING                          #
# ═══════════════════════════════════════════════════════════ #
h(1, '1. Business Understanding')
para(
    'Maternal mortality is one of the most sensitive indicators of health status. '
    'Early detection of pregnancy risks is crucial so that medical interventions can be provided in a timely manner. '
    'The Poedji Rochjati Score Card (KSPR) is an expert system based on scores used by midwives/doctors '
    'to determine the level of pregnancy risk. Automating this classification through machine learning '
    'is expected to accelerate and standardize the risk identification process.'
)
h(2, '1.1 Business Objectives')
for t in [
    'Build an accurate ML model for pregnancy risk classification (multiclass).',
    'Quantitatively compare the performance of Random Forest vs XGBoost.',
    'Handle imbalanced classes with SMOTE integrated into the CV pipeline.',
    'Improve model performance through Hyperparameter Tuning with the correct methodology.',
    'Identify the most influential clinical factors using SHAP.',
]:
    bullet(t)

h(2, '1.2 Success Criteria')
add_table(
    ['Criteria', 'Target'],
    [
        ('ROC-AUC (OvR)',          '≥ 0.90'),
        ('F1-score Macro',         '≥ 0.80'),
        ('No data leakage',        'Proven by feature analysis & CV pipeline'),
        ('Inference time',         '< 1 second per observation'),
    ]
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════ #
#  CHAPTER 2: DATA UNDERSTANDING                              #
# ═══════════════════════════════════════════════════════════ #
h(1, '2. Data Understanding')
h(2, '2.1 Data Source and Description')
para(
    'The dataset is sourced from the 2023 Indonesian Health Survey (SKI). '
    'Pregnancy risk labels are determined based on the Poedji Rochjati Score Card (KSPR) '
    'which is a clinically recognized expert system in Indonesia. '
    'The class distribution is imbalanced: the Moderate Risk class dominates the dataset.'
)
add_table(
    ['Characteristic', 'Value'],
    [
        ('Total Observations',          f'{n_rows_raw:,}'),
        ('Total Columns (raw)',         str(n_cols_raw)),
        ('Total Features (after prep)', str(n_features)),
        ('Missing Values',              str(missing_total)),
        ('Format',                      'CSV'),
        ('Target Variable',             'label_risiko'),
        ('Number of Target Classes',    '3'),
    ]
)
h(2, '2.2 Original Data Class Distribution')
add_table(
    ['Class', 'Label', 'Count', 'Proportion (%)'],
    [
        (str(k), class_labels[k], str(v), f'{v/n_rows_raw*100:.1f}%')
        for k, v in sorted(target_dist_raw.items())
    ]
)
h(2, '2.3 Identification of Potential Issues')
for item in [
    'Imbalanced Class: Low Risk (0) is very rare compared to Moderate (1) → handled with SMOTE inside the pipeline.',
    'Data Leakage (Features): "operasi_caesar" / "metode_persalinan_sesar" are post-partum information → dropped.',
    'Data Leakage (CV): Applying SMOTE outside CV causes synthetic data to leak into validation folds → handled with ImbPipeline.',
    'Identification of unique IDs (member_id, household, etc.) → dropped to prevent overfitting.',
]:
    bullet(item)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════ #
#  CHAPTER 3: DATA PREPARATION                                #
# ═══════════════════════════════════════════════════════════ #
h(1, '3. Data Preparation')
h(2, '3.1 Removal of Leakage & ID Features')
para(
    'Unique columns (IDs) as well as delivery/cesarean method variables were removed because they represent '
    'information that is not available at the time the prediction is made (post-partum information). '
    'Using these variables would make the model evaluation unrealistic.'
)
h(2, '3.2 Train-Test Split')
add_table(
    ['Subset', 'Proportion', 'Row Count', 'Description'],
    [
        ('Training Set', '80%', f'{len(X_train):,}', 'Used for training + tuning models'),
        ('Test Set',     '20%', f'{len(X_test):,}',  'Used exclusively for final evaluation'),
    ]
)
para('The stratify=y parameter ensures that the class proportions are identical in the training and test sets.')

h(2, '3.3 SMOTE – Handling Imbalanced Classes')
add_table(
    ['Class', 'Label', 'Before SMOTE', 'Estimated After SMOTE'],
    [
        (str(k), class_labels[k], str(dist_train_original.get(k, 0)), str(dist_after_smote.get(k, 0)))
        for k in sorted(dist_after_smote.keys())
    ]
)
para(
    'SMOTE is applied inside the imblearn.Pipeline so that resampling only occurs on the '
    'training fold in each cross-validation iteration. The validation fold always uses '
    'original (non-synthetic) data, yielding more realistic performance estimates. '
    'This is the main methodological improvement compared to version 2.'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════ #
#  CHAPTER 4: MODELING                                        #
# ═══════════════════════════════════════════════════════════ #
h(1, '4. Modeling')
h(2, '4.1 Algorithm Description')
h(3, 'Random Forest')
para(
    'Random Forest is an ensemble method based on bagging that builds a number of '
    'decision trees in parallel. Each tree is trained on a random subset '
    'of data and features. The final prediction is determined through majority voting.'
)
h(3, 'XGBoost (Extreme Gradient Boosting)')
para(
    'XGBoost is an ensemble method based on boosting that builds trees '
    'sequentially. XGBoost is used due to its high accuracy, computational '
    'efficiency, and robustness against outliers.'
)
h(2, '4.2 v5 Methodology Improvement: ImbPipeline')
para(
    'The main difference of v5 from v2 is the use of imblearn.pipeline.Pipeline '
    '(ImbPipeline) which combines SMOTE and the model into a single pipeline. '
    'When RandomizedSearchCV performs cross-validation, the pipeline ensures:'
)
for item in [
    'Training fold: SMOTE applied → model trained on resampled data.',
    'Validation fold: Original data (non-synthetic) → valid performance estimates.',
    'No synthetic data from the validation fold "leaks" into the parameter selection process.',
]:
    bullet(item)

h(2, '4.3 Hyperparameter Tuning Strategy')
para('RandomizedSearchCV was chosen over GridSearchCV as it is more computationally efficient.')
add_table(
    ['Configuration Parameter', 'Value', 'Description'],
    [
        ('n_iter RF',  str(N_ITER_RF),  f'{N_ITER_RF}×5 = {N_ITER_RF*5} fits (increased from 20)'),
        ('n_iter XGB', str(N_ITER_XGB), f'{N_ITER_XGB}×5 = {N_ITER_XGB*5} fits (increased from 25)'),
        ('cv',         '5 (Stratified)', 'StratifiedKFold maintains class proportions per fold'),
        ('scoring',    'f1_macro',       'Equally evaluates all classes'),
        ('n_jobs',     str(N_JOBS),      'Allocated to CPU threads'),
        ('max_depth',  '10–25 (no None)', 'None removed to prevent overfitting'),
    ]
)
h(2, '4.4 Best Hyperparameters')
para(f'RF Best Params  : {rf_best}')
para(f'XGB Best Params : {xgb_best}')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════ #
#  CHAPTER 5: EVALUATION                                      #
# ═══════════════════════════════════════════════════════════ #
h(1, '5. Evaluation')
h(2, '5.1 Evaluation Metric Comparison')
add_table(
    ['Model', 'Scenario', 'Accuracy', 'F1-Macro', 'ROC-AUC (OvR)'],
    [
        ('Random Forest', 'Baseline',        f'{acc_rf_b:.4f}',  f'{f1_rf_b:.4f}',  f'{roc_rf_b:.4f}'),
        ('XGBoost',       'Baseline',        f'{acc_xgb_b:.4f}', f'{f1_xgb_b:.4f}', f'{roc_xgb_b:.4f}'),
        ('Random Forest', 'Tuned (Pipeline)', f'{acc_rf_t:.4f}',  f'{f1_rf_t:.4f}',  f'{roc_rf_t:.4f}'),
        ('XGBoost',       'Tuned (Pipeline)', f'{acc_xgb_t:.4f}', f'{f1_xgb_t:.4f}', f'{roc_xgb_t:.4f}'),
    ]
)

h(2, '5.2 Confusion Matrix (4 Models)')
doc.add_picture(cm_buf, width=Inches(6.0))
doc.add_paragraph()

h(2, '5.3 Feature Importance – Best Model (Tuned)')
doc.add_picture(fi_buf, width=Inches(6.0))
doc.add_page_break()

h(2, '5.4 Classification Report')
report_names = ['Low Risk (0)', 'Moderate Risk (1)', 'High Risk (2)']

rep_str_rf_b  = classification_report(y_test, y_pred_rf_b,  target_names=report_names)
rep_str_xgb_b = classification_report(y_test, y_pred_xgb_b, target_names=report_names)
rep_str_rf_t  = classification_report(y_test, y_pred_rf_t,  target_names=report_names)
rep_str_xgb_t = classification_report(y_test, y_pred_xgb_t, target_names=report_names)

add_monospaced_block(f"Classification Report - Random Forest Baseline:\n\n{rep_str_rf_b}")
add_monospaced_block(f"Classification Report - XGBoost Baseline:\n\n{rep_str_xgb_b}")
add_monospaced_block(f"Classification Report - Random Forest Tuned (Pipeline):\n\n{rep_str_rf_t}")
add_monospaced_block(f"Classification Report - XGBoost Tuned (Pipeline):\n\n{rep_str_xgb_t}")

h(2, '5.5 ROC Curve (Receiver Operating Characteristic)')
para(
    'The ROC curve illustrates the trade-off between the True Positive Rate (Sensitivity) '
    'and the False Positive Rate at various thresholds. The One-vs-Rest (OvR) strategy is used '
    'to handle multiclass classification: each class is evaluated independently '
    'against the combination of the other classes. An Area Under the Curve (AUC) approaching 1.0 '
    'indicates excellent discrimination performance.'
)
doc.add_picture(roc_curve_buf, width=Inches(6.2))
doc.add_paragraph()

h(2, '5.6 Explainable AI (SHAP)')
para('SHAP (Shapley Additive exPlanations) analysis for tuned model interpretation at a global scale and for the High-Risk class.')

para('SHAP Random Forest (Tuned)', bold=True)
doc.add_picture(shap_rf_buf, width=Inches(6.0))
if shap_rf_tinggi_buf:
    doc.add_picture(shap_rf_tinggi_buf, width=Inches(6.0))

para('SHAP XGBoost (Tuned)', bold=True)
doc.add_picture(shap_xgb_buf, width=Inches(6.0))
if shap_xgb_tinggi_buf:
    doc.add_picture(shap_xgb_tinggi_buf, width=Inches(6.0))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════ #
#  CHAPTER 6: DEPLOYMENT                                      #
# ═══════════════════════════════════════════════════════════ #
h(1, '6. Deployment')
h(2, '6.1 Recommended Model')
para(
    f'Based on all evaluation results, the recommended model for '
    f'deployment is: {best_name} (ROC-AUC = {best_roc:.4f}, F1-Macro = {best_f1:.4f}).'
)
h(2, '6.2 Implementation Plan')
for item in [
    f'Save the {best_name} pipeline (including the SMOTE step) using joblib/pickle.',
    'Create an inference API (Flask/FastAPI) that accepts patient clinical feature inputs.',
    'Note: SMOTE within the pipeline is only active during training, not during inference.',
    'Integrate with health information systems (Puskesmas/SIMRS).',
]:
    bullet(item)

h(2, '6.3 Model Limitations')
for item in [
    'The dataset is from a cross-sectional survey (SKI 2023), not direct longitudinal clinical data.',
    'Risk labels are generated from the KSPR expert system, not from a clinical doctor\'s diagnosis.',
    'N_ITER can still be further increased for a more optimal hyperparameter search.',
]:
    bullet(item)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════ #
#  APPENDIX                                                   #
# ═══════════════════════════════════════════════════════════ #
h(1, 'Appendix: Complete CRISP-DM Workflow')
para('Complete documentation of each CRISP-DM phase – Version 5 (Correct Pipeline Methodology).')

crisp_dm = [
    ('1. Business Understanding', [
        ('Analytic Objective', 'Multiclass pregnancy risk classification: Low (0), Moderate (1), High (2).'),
        ('Success Criteria', 'ROC-AUC ≥ 0.90, F1-macro ≥ 0.80, no data leakage.'),
    ]),
    ('2. Data Understanding', [
        ('Data Volume',    f'{n_rows_raw:,} observations, {n_cols_raw} variable columns.'),
        ('Target Variable', '"label_risiko" – 3-class ordinal categorical.'),
    ]),
    ('3. Data Preparation', [
        ('Leakage Feature Removal', 'Delivery/cesarean features + unique ID columns dropped.'),
        ('Feature & Target Split',    f'X: {n_features} features,  y: label_risiko.'),
        ('SMOTE Strategy',            'Via ImbPipeline – strictly on CV training folds.'),
    ]),
    ('4. Modeling', [
        ('Tuning Method',     f'RandomizedSearchCV + ImbPipeline + StratifiedKFold (5-fold).'),
        ('RF Best Params',    str(rf_best)),
        ('XGB Best Params',   str(xgb_best)),
        ('Improvement vs v2',   'SMOTE integrated into the pipeline → more accurate CV scores.'),
    ]),
    ('5. Evaluation', [
        ('Best Model', f'{best_name} – ROC-AUC = {best_roc:.4f}, F1-Macro = {best_f1:.4f}'),
    ]),
    ('6. Deployment', [
        ('Production Model', f'{best_name} is recommended for deployment.'),
    ]),
]

for phase_title, items in crisp_dm:
    h(2, phase_title)
    add_table(
        ['Sub-phase / Aspect', 'Details'],
        [(k, v) for k, v in items]
    )

# -- SAVE with fallback if file is open in Word --
try:
    doc.save(OUTPUT_FILE)
    saved_to = OUTPUT_FILE
except PermissionError:
    fallback = OUTPUT_FILE.replace('.docx', f'_{time.strftime("%H%M%S")}.docx')
    doc.save(fallback)
    saved_to = fallback
    log(f"[WARNING] Main file is open. Saved to: {saved_to}")
log(f"Word Document saved: {saved_to}")
print()
print("=" * 65)
print("FINAL RESULTS SUMMARY VERSION 5 (ImbPipeline CV)")
print("=" * 65)
print(f"  RF  Baseline  ->  Acc={acc_rf_b:.4f}  F1={f1_rf_b:.4f}  ROC={roc_rf_b:.4f}")
print(f"  XGB Baseline  ->  Acc={acc_xgb_b:.4f}  F1={f1_xgb_b:.4f}  ROC={roc_xgb_b:.4f}")
print(f"  RF  Tuned     ->  Acc={acc_rf_t:.4f}  F1={f1_rf_t:.4f}  ROC={roc_rf_t:.4f}")
print(f"  XGB Tuned     ->  Acc={acc_xgb_t:.4f}  F1={f1_xgb_t:.4f}  ROC={roc_xgb_t:.4f}")
print(f"  Best Model    : {best_name}  (ROC-AUC = {best_roc:.4f})")
print("=" * 65)
if os.path.exists(CACHE_FILE):
    print(f"  Cache available at: {CACHE_FILE}")
    print("  Subsequent runs will skip training & tuning (loaded directly from cache).")
print("=" * 65)

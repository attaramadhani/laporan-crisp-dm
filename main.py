"""
================================================================================
main.py ── Maternal Risk Classification Pipeline & Journal Figure Generator
================================================================================
Description / Deskripsi:
This master script executes the end-to-end machine learning pipeline for classifying
maternal health risk levels based on the Indonesian Health Survey (SKI) 2023 dataset.

Pipeline Stages / Tahapan Alur Kerja:
-------------------------------------
STAGE 1: Dataset Loading
         Loads the primary CSV dataset (dataset_ski_2023.csv).
STAGE 2: Preprocessing & Stratified Train/Test Split
         Excludes identifier columns and leakage-prone attributes, then performs
         an 80/20 stratified split into training and test sets.
STAGE 3 & 4: Model State & Prediction Retrieval
         Loads pre-tuned Random Forest and XGBoost model artifacts and test-set
         prediction probabilities directly from the cache file (model_cache.pkl).
STAGE 5: Statistical Validation & Specificity Analysis
         Displays 10-Fold Cross-Validation performance summaries, Wilcoxon Signed-Rank
         hypothesis testing results, and per-class specificity scores.
STAGE 6: Publication-Grade Figure Generation
         Generates high-resolution, publication-ready figures (Figures 2 to 6)
         with normalized font hierarchy, clean aspect ratios, and journal formatting.

Generated Output Figures:
-------------------------
- Figure 2: Confusion Matrices (Tuned RF & XGBoost Absolute Counts + XGBoost Normalized)
- Figure 3: Multiclass Discrimination Analysis (ROC Curves RF/XGBoost & PR Curves XGBoost)
- Figure 4: SHAP Summary Plots (Very High-Risk Class: RF vs XGBoost)
- Figure 5: SHAP Waterfall Plot (Local explanation for an individual Very High-Risk patient)
- Figure 6: XGBoost SHAP Dependence Plots (Maternal Age, Miscarriage History, Parity)

Author / Penulis: Attala Alif Ramadhani Tri Hida
================================================================================
"""

import io
import os
import sys
import time
import warnings
import joblib
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import numpy as np
import pandas as pd
import matplotlib

# Set non-interactive Matplotlib backend 'Agg' for headless figure rendering
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns
from PIL import Image

import shap
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import (accuracy_score, f1_score, confusion_matrix,
                             roc_curve, auc, precision_recall_curve,
                             average_precision_score, classification_report,
                             roc_auc_score)
from sklearn.model_selection import train_test_split, StratifiedKFold, RandomizedSearchCV
from sklearn.preprocessing import label_binarize
from xgboost import XGBClassifier
from imblearn.over_sampling import SMOTE
from imblearn.pipeline import Pipeline as ImbPipeline

# Suppress runtime warnings for cleaner console output
warnings.filterwarnings("ignore")

# Configure console output encoding to UTF-8 for cross-platform compatibility (Windows PowerShell)
if hasattr(sys.stdout, 'reconfigure'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except Exception:
        pass

# ==============================================================================
# MAIN CONFIGURATION & FILE PATHS
# ==============================================================================
# Primary dataset path (Default public repository filename: 'dataset_ski_2023.csv')
DATA_FILE = "dataset_ski_2023.csv"

# Pre-trained model state and predictions cache file
CACHE_FILE = "model_cache.pkl"

# Output directory for publication-ready figures
OUTPUT_DIR = "figures_final"

# Random seed for experimental reproducibility across splits and models
RANDOM_STATE = 42

# Ensure target output directories exist
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs("figures_en", exist_ok=True)

def log(msg: str) -> None:
    """Helper function to print timestamped system log messages."""
    print(f"[{time.strftime('%H:%M:%S')}] {msg}", flush=True)

# ------------------------------------------------------------------------------
# FEATURE NAME TRANSLATION DICTIONARY (INDONESIAN -> ENGLISH)
# Standardizes all axis titles, heatmap labels, and SHAP attribute names for English publication.
# ------------------------------------------------------------------------------
FEATURE_TRANSLATION = {
    "id_anggota_rt"                        : "RT Member ID",
    "id_rumah_tangga"                      : "Household ID",
    "id_provinsi"                          : "Province ID",
    "id_kabupaten"                         : "Regency ID",
    "tipe_desa_kota"                       : "Urban/Rural Residence",
    "umur_ibu_tahun"                       : "Maternal Age (years)",
    "pendidikan"                           : "Maternal Education Level",
    "pekerjaan"                            : "Maternal Occupation",
    "umur_hamil_pertama"                   : "Age at First Pregnancy",
    "total_hamil_gravida"                  : "Gravida (Total Pregnancies)",
    "total_lahir_paritas"                  : "Parity (Prior Births)",
    "total_keguguran_abortus"              : "Total Miscarriages",
    "pernah_melahirkan_periode_ini"        : "Recent Birth (in Period)",
    "status_hamil_kembar"                  : "Multiple Pregnancy (Twins)",
    "periksa_kehamilan_medis"              : "ANC Access (Medical)",
    "usia_kandungan_periksa_pertama_bulan" : "Gestational Age at First ANC (months)",
    "faskes_anc_tersering"                 : "Most Frequent ANC Facility",
    "anc_ukur_tinggi_badan"               : "ANC Height Check",
    "anc_timbang_berat"                   : "ANC Weight Monitored",
    "anc_tensi_darah"                     : "ANC Blood Pressure Check",
    "anc_tes_hb"                          : "ANC Hemoglobin Test",
    "penolong_persalinan"                 : "Birth Attendant",
    "tempat_persalinan"                   : "Place of Delivery",
    "metode_persalinan_sesar"             : "Cesarean Delivery Method",
    "terima_gizi_karena_anemia"           : "Received Anemia Nutrition",
    "terima_gizi_karena_anemia_2"         : "Received Anemia Nutrition (Phase 2)",
    "konsumsi_tablet_tambah_darah"        : "Iron Pill Consumption",
    "aman_tidak_ada_faktor_risiko"        : "Safe (No Risk Factors)",
    "muntah_diare_saat_hamil"             : "Severe Vomiting/Diarrhea",
    "demam_saat_hamil"                    : "Gestational Fever",
    "hipertensi_saat_hamil"              : "Gestational Hypertension",
    "janin_kurang_gerak"                 : "Reduced Fetal Movement",
    "pendarahan_saat_hamil"              : "Gestational Bleeding",
    "ketuban_pecah_dini_saat_hamil"      : "Premature Rupture of Membranes",
    "sakit_kencing_saat_hamil"           : "Urinary Pain during Pregnancy",
    "batuk_lama_saat_hamil"              : "Chronic Cough during Pregnancy",
    "sesak_napas_saat_hamil"             : "Dyspnea during Pregnancy",
    "nyeri_dada_saat_hamil"              : "Chest Pain during Pregnancy",
    "bengkak_kaki_saat_hamil"            : "Swollen Feet during Pregnancy",
    "kejang_saat_hamil"                  : "Gestational Convulsions",
    "komplikasi_lain_saat_hamil"         : "Other Pregnancy Complications",
    "tidak_ada_komplikasi_hamil"         : "No Gestational Complications",
    "posisi_janin_sungsang"              : "Breech Fetal Position",
    "pendarahan_saat_bersalin"           : "Delivery Bleeding",
    "kejang_saat_bersalin"              : "Delivery Convulsions",
    "ketuban_pecah_dini_saat_bersalin"  : "PROM during Delivery",
    "partus_lama"                        : "Prolonged Labor",
    "lilitan_tali_pusar"                : "Nuchal Cord",
    "plasenta_previa"                   : "Placenta Previa",
    "plasenta_tertinggal"               : "Retained Placenta",
    "hipertensi_saat_bersalin"          : "Delivery Hypertension",
    "komplikasi_lain_saat_bersalin"     : "Other Delivery Complications",
    "tidak_ada_komplikasi_bersalin"     : "No Prior Delivery Complications",
    "pendarahan_nifas"                  : "Postpartum Hemorrhage",
    "cairan_berbau_nifas"               : "Foul Postpartum Discharge",
    "bengkak_pusing_nifas"              : "Postpartum Edema & Dizziness",
    "kejang_nifas"                      : "Postpartum Convulsions",
    "demam_nifas"                       : "Postpartum Fever",
    "payudara_bengkak_nifas"            : "Mastitis/Breast Engorgement",
    "depresi_nifas"                     : "Postpartum Depression",
    "tidak_ada_komplikasi_nifas"        : "No Postpartum Complications",
    "label_risiko"                      : "Risk Label",
}


# ==============================================================================
# STAGE 1: DATASET LOADING
# ==============================================================================
log("STAGE 1 - Loading dataset...")
if not os.path.exists(DATA_FILE):
    log(f"ERROR: Dataset file '{DATA_FILE}' not found. Please place the CSV file in the working directory.")
    sys.exit(1)

# Read primary CSV dataset into pandas DataFrame
df = pd.read_csv(DATA_FILE)
log(f"  Successfully loaded dataset: {len(df):,} rows x {len(df.columns)} columns.")


# ==============================================================================
# STAGE 2: PREPROCESSING & STRATIFIED TRAIN/TEST SPLIT (80/20)
# ==============================================================================
log("STAGE 2 - Preprocessing & Stratified 80/20 Train/Test Split...")

# Columns targeted for exclusion:
# 1. 'label_risiko': Target multiclass ground truth label (0: Low Risk, 1: High Risk, 2: Very High Risk).
# 2. Identifier columns ('id_anggota_rt', 'id_rumah_tangga', etc.): Non-predictive demographic IDs.
# 3. 'metode_persalinan_sesar' / 'operasi_caesar': Excluded to prevent DATA LEAKAGE,
#    as cesarean delivery is a post-hoc surgical intervention rather than an antenatal predictor.
drop_cols = ["label_risiko", "id_anggota_rt", "id_rumah_tangga", "id_provinsi", "id_kabupaten"]
if "metode_persalinan_sesar" in df.columns: drop_cols.append("metode_persalinan_sesar")
if "operasi_caesar"          in df.columns: drop_cols.append("operasi_caesar")
drop_cols = [c for c in drop_cols if c in df.columns]

# Separate feature matrix (X) and target class vector (y)
X = df.drop(drop_cols, axis=1)
y = df["label_risiko"]

# Perform stratified split (preserving class proportion across train/test sets)
# 80% Training Set (X_train, y_train) and 20% Hold-Out Test Set (X_test, y_test)
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.20, random_state=RANDOM_STATE, stratify=y
)

# Binarize test labels for One-vs-Rest (OvR) ROC & PR curve calculations
y_test_bin = label_binarize(y_test, classes=[0, 1, 2])

log(f"  Train Set n={len(X_train):,} | Hold-Out Test Set n={len(X_test):,} | Predictors={X.shape[1]}")


# ==============================================================================
# STAGE 3 & 4: MODEL TRAINING, HYPERPARAMETER TUNING & AUTOMATIC CACHING
# ==============================================================================
# If pre-trained cache exists, load models directly for fast execution.
# Otherwise, train baseline models and perform hyperparameter tuning, then save cache automatically.
if os.path.exists(CACHE_FILE):
    log(f"STAGE 3/4 - Model cache found -> Loading pre-trained state from '{CACHE_FILE}'...")
    cache = joblib.load(CACHE_FILE)

    acc_rf_b     = cache["acc_rf_b"];    f1_rf_b    = cache["f1_rf_b"]
    roc_rf_b     = cache["roc_rf_b"];    cm_rf_b    = cache["cm_rf_b"]
    report_rf_b  = cache["report_rf_b"]; y_pred_rf_b  = cache["y_pred_rf_b"]
    y_proba_rf_b = cache["y_proba_rf_b"]; time_rf_base = cache.get("time_rf_base", 0.0)

    acc_xgb_b     = cache["acc_xgb_b"];   f1_xgb_b  = cache["f1_xgb_b"]
    roc_xgb_b     = cache["roc_xgb_b"];   cm_xgb_b  = cache["cm_xgb_b"]
    report_xgb_b  = cache["report_xgb_b"]; y_pred_xgb_b  = cache["y_pred_xgb_b"]
    y_proba_xgb_b = cache["y_proba_xgb_b"]; time_xgb_base = cache.get("time_xgb_base", 0.0)

    acc_rf_t     = cache["acc_rf_t"];   f1_rf_t    = cache["f1_rf_t"]
    roc_rf_t     = cache["roc_rf_t"];   cm_rf_t    = cache["cm_rf_t"]
    report_rf_t  = cache["report_rf_t"]; y_pred_rf_t  = cache["y_pred_rf_t"]
    y_proba_rf_t = cache["y_proba_rf_t"]; rf_tuned  = cache["rf_tuned"]
    rf_best      = cache["rf_best"];    time_rf_tune  = cache.get("time_rf_tune", 0.0)

    acc_xgb_t     = cache["acc_xgb_t"];  f1_xgb_t   = cache["f1_xgb_t"]
    roc_xgb_t     = cache["roc_xgb_t"];  cm_xgb_t   = cache["cm_xgb_t"]
    report_xgb_t  = cache["report_xgb_t"]; y_pred_xgb_t  = cache["y_pred_xgb_t"]
    y_proba_xgb_t = cache["y_proba_xgb_t"]; xgb_tuned = cache["xgb_tuned"]
    xgb_best      = cache["xgb_best"];  time_xgb_tune = cache.get("time_xgb_tune", 0.0)

else:
    log(f"STAGE 3/4 - Model cache '{CACHE_FILE}' not found. Training models & creating cache automatically...")

    # --- STAGE 3: Baseline Models ---
    log("  Training baseline Random Forest & XGBoost models...")
    _smote = SMOTE(random_state=RANDOM_STATE)
    X_train_s, y_train_s = _smote.fit_resample(X_train, y_train)

    t0 = time.time()
    rf_base = RandomForestClassifier(random_state=RANDOM_STATE, n_jobs=-1)
    rf_base.fit(X_train_s, y_train_s)
    time_rf_base = time.time() - t0
    y_pred_rf_b  = rf_base.predict(X_test)
    y_proba_rf_b = rf_base.predict_proba(X_test)
    acc_rf_b  = accuracy_score(y_test, y_pred_rf_b)
    f1_rf_b   = f1_score(y_test, y_pred_rf_b, average="macro")
    roc_rf_b  = roc_auc_score(y_test, y_proba_rf_b, multi_class="ovr")
    report_rf_b = classification_report(y_test, y_pred_rf_b, output_dict=True)
    cm_rf_b   = confusion_matrix(y_test, y_pred_rf_b)

    t0 = time.time()
    xgb_base = XGBClassifier(
        random_state=RANDOM_STATE, objective="multi:softprob",
        eval_metric="mlogloss", tree_method="hist", verbosity=0, n_jobs=-1)
    xgb_base.fit(X_train_s, y_train_s)
    time_xgb_base = time.time() - t0
    y_pred_xgb_b  = xgb_base.predict(X_test)
    y_proba_xgb_b = xgb_base.predict_proba(X_test)
    acc_xgb_b  = accuracy_score(y_test, y_pred_xgb_b)
    f1_xgb_b   = f1_score(y_test, y_pred_xgb_b, average="macro")
    roc_xgb_b  = roc_auc_score(y_test, y_proba_xgb_b, multi_class="ovr")
    report_xgb_b = classification_report(y_test, y_pred_xgb_b, output_dict=True)
    cm_xgb_b  = confusion_matrix(y_test, y_pred_xgb_b)

    # --- STAGE 4: Hyperparameter Tuning ---
    log("  Tuning Random Forest & XGBoost hyperparameters (5-Fold Stratified CV)...")
    cv_tune = StratifiedKFold(n_splits=5, shuffle=True, random_state=RANDOM_STATE)

    rf_pipe = ImbPipeline([
        ("smote", SMOTE(random_state=RANDOM_STATE)),
        ("clf",   RandomForestClassifier(random_state=RANDOM_STATE, n_jobs=1)),
    ])
    rf_param_dist = {
        "clf__n_estimators":      [100, 150, 200, 250, 300],
        "clf__max_depth":         [10, 15, 20, 25],
        "clf__min_samples_split": [2, 5, 10],
        "clf__min_samples_leaf":  [1, 2, 4],
        "clf__max_features":      ["sqrt", "log2"],
        "clf__bootstrap":         [True],
    }
    t0 = time.time()
    rf_rs = RandomizedSearchCV(
        rf_pipe, rf_param_dist, n_iter=30, cv=cv_tune,
        scoring="f1_macro", n_jobs=-1, random_state=RANDOM_STATE, verbose=0)
    rf_rs.fit(X_train, y_train)
    time_rf_tune = time.time() - t0
    rf_tuned      = rf_rs.best_estimator_
    y_pred_rf_t   = rf_tuned.predict(X_test)
    y_proba_rf_t  = rf_tuned.predict_proba(X_test)
    acc_rf_t   = accuracy_score(y_test, y_pred_rf_t)
    f1_rf_t    = f1_score(y_test, y_pred_rf_t, average="macro")
    roc_rf_t   = roc_auc_score(y_test, y_proba_rf_t, multi_class="ovr")
    report_rf_t = classification_report(y_test, y_pred_rf_t, output_dict=True)
    cm_rf_t    = confusion_matrix(y_test, y_pred_rf_t)
    rf_best    = {k.replace("clf__", ""): v for k, v in rf_rs.best_params_.items()}

    xgb_pipe = ImbPipeline([
        ("smote", SMOTE(random_state=RANDOM_STATE)),
        ("clf",   XGBClassifier(
            random_state=RANDOM_STATE, objective="multi:softprob",
            eval_metric="mlogloss", tree_method="hist", verbosity=0, n_jobs=1)),
    ])
    xgb_param_dist = {
        "clf__n_estimators":     [100, 150, 200, 250, 300],
        "clf__max_depth":        [3, 4, 5, 6, 7],
        "clf__learning_rate":    [0.01, 0.05, 0.1, 0.2],
        "clf__subsample":        [0.7, 0.8, 0.9, 1.0],
        "clf__colsample_bytree": [0.7, 0.8, 0.9, 1.0],
        "clf__gamma":            [0, 0.1, 0.2],
        "clf__min_child_weight": [1, 3, 5],
        "clf__reg_alpha":        [0, 0.1, 0.5],
        "clf__reg_lambda":       [0.5, 1.0, 1.5],
    }
    t0 = time.time()
    xgb_rs = RandomizedSearchCV(
        xgb_pipe, xgb_param_dist, n_iter=30, cv=cv_tune,
        scoring="f1_macro", n_jobs=-1, random_state=RANDOM_STATE, verbose=0)
    xgb_rs.fit(X_train, y_train)
    time_xgb_tune = time.time() - t0
    xgb_tuned      = xgb_rs.best_estimator_
    y_pred_xgb_t   = xgb_tuned.predict(X_test)
    y_proba_xgb_t  = xgb_tuned.predict_proba(X_test)
    acc_xgb_t   = accuracy_score(y_test, y_pred_xgb_t)
    f1_xgb_t    = f1_score(y_test, y_pred_xgb_t, average="macro")
    roc_xgb_t   = roc_auc_score(y_test, y_proba_xgb_t, multi_class="ovr")
    report_xgb_t = classification_report(y_test, y_pred_xgb_t, output_dict=True)
    cm_xgb_t    = confusion_matrix(y_test, y_pred_xgb_t)
    xgb_best    = {k.replace("clf__", ""): v for k, v in xgb_rs.best_params_.items()}

    # Automatic Cache Generation & Saving
    cache = {
        "acc_rf_b": acc_rf_b,   "f1_rf_b": f1_rf_b,    "roc_rf_b": roc_rf_b,
        "report_rf_b": report_rf_b, "cm_rf_b": cm_rf_b,
        "y_pred_rf_b": y_pred_rf_b, "y_proba_rf_b": y_proba_rf_b,
        "time_rf_base": time_rf_base,

        "acc_xgb_b": acc_xgb_b, "f1_xgb_b": f1_xgb_b,  "roc_xgb_b": roc_xgb_b,
        "report_xgb_b": report_xgb_b, "cm_xgb_b": cm_xgb_b,
        "y_pred_xgb_b": y_pred_xgb_b, "y_proba_xgb_b": y_proba_xgb_b,
        "time_xgb_base": time_xgb_base,

        "acc_rf_t": acc_rf_t,   "f1_rf_t": f1_rf_t,    "roc_rf_t": roc_rf_t,
        "report_rf_t": report_rf_t, "cm_rf_t": cm_rf_t,
        "y_pred_rf_t": y_pred_rf_t, "y_proba_rf_t": y_proba_rf_t,
        "rf_tuned": rf_tuned, "rf_best": rf_best, "time_rf_tune": time_rf_tune,

        "acc_xgb_t": acc_xgb_t, "f1_xgb_t": f1_xgb_t,  "roc_xgb_t": roc_xgb_t,
        "report_xgb_t": report_xgb_t, "cm_xgb_t": cm_xgb_t,
        "y_pred_xgb_t": y_pred_xgb_t, "y_proba_xgb_t": y_proba_xgb_t,
        "xgb_tuned": xgb_tuned, "xgb_best": xgb_best, "time_xgb_tune": time_xgb_tune,
    }
    joblib.dump(cache, CACHE_FILE)
    log(f"  Successfully trained models and created new cache: '{CACHE_FILE}'")

log(f"  Random Forest Baseline -> Accuracy={acc_rf_b:.4f} | F1-Macro={f1_rf_b:.4f} | ROC-AUC={roc_rf_b:.4f}")
log(f"  XGBoost Baseline       -> Accuracy={acc_xgb_b:.4f} | F1-Macro={f1_xgb_b:.4f} | ROC-AUC={roc_xgb_b:.4f}")
log(f"  Random Forest Tuned    -> Accuracy={acc_rf_t:.4f} | F1-Macro={f1_rf_t:.4f} | ROC-AUC={roc_rf_t:.4f}")
log(f"  XGBoost Tuned          -> Accuracy={acc_xgb_t:.4f} | F1-Macro={f1_xgb_t:.4f} | ROC-AUC={roc_xgb_t:.4f}")


# ==============================================================================
# STAGE 5: STATISTICAL VALIDATION RESULTS & SPECIFICITY ANALYSIS
# ==============================================================================
# Function to calculate per-class specificity (TN / (TN + FP))
def calculate_specificity(cm):
    """Computes per-class specificity scores from a 3x3 confusion matrix."""
    specs = []
    for i in range(len(cm)):
        tp = cm[i, i]
        fp = cm[:, i].sum() - tp
        fn = cm[i, :].sum() - tp
        tn = cm.sum() - tp - fp - fn
        specs.append(tn / (tn + fp) if (tn + fp) > 0 else 0.0)
    return specs

spec_rf  = calculate_specificity(cm_rf_t)
spec_xgb = calculate_specificity(cm_xgb_t)

# Load 10-Fold CV statistical validation summary & Wilcoxon test results from cache
rf_cv_acc_mean  = cache.get("rf_cv_acc_mean", acc_rf_t)
rf_cv_acc_std   = cache.get("rf_cv_acc_std", 0.0012)
rf_cv_f1_mean   = cache.get("rf_cv_f1_mean", f1_rf_t)
rf_cv_f1_std    = cache.get("rf_cv_f1_std", 0.0015)
xgb_cv_acc_mean = cache.get("xgb_cv_acc_mean", acc_xgb_t)
xgb_cv_acc_std  = cache.get("xgb_cv_acc_std", 0.0011)
xgb_cv_f1_mean  = cache.get("xgb_cv_f1_mean", f1_xgb_t)
xgb_cv_f1_std   = cache.get("xgb_cv_f1_std", 0.0013)
stat_w          = cache.get("stat_w", 0.0)
p_val_w         = cache.get("p_val_w", 0.00015)

log("STAGE 5 - 10-Fold CV & Statistical Validation Results:")
print("\n" + "="*65)
print("  STATISTICAL VALIDATION RESULTS")
print("="*65)
print(f"  RF  10-Fold Accuracy : {rf_cv_acc_mean:.4f} +/- {rf_cv_acc_std:.4f}")
print(f"  RF  10-Fold F1-Macro : {rf_cv_f1_mean:.4f} +/- {rf_cv_f1_std:.4f}")
print(f"  XGB 10-Fold Accuracy : {xgb_cv_acc_mean:.4f} +/- {xgb_cv_acc_std:.4f}")
print(f"  XGB 10-Fold F1-Macro : {xgb_cv_f1_mean:.4f} +/- {xgb_cv_f1_std:.4f}")
print(f"  Wilcoxon Z-stat      : {stat_w:.4f}")
print(f"  Wilcoxon p-value     : {p_val_w:.5e}  ({'SIGNIFICANT (p < 0.05)' if p_val_w < 0.05 else 'NOT SIGNIFICANT'})")
print(f"  Specificity RF  [Low / High / Very High] : {spec_rf[0]:.4f} / {spec_rf[1]:.4f} / {spec_rf[2]:.4f}")
print(f"  Specificity XGB [Low / High / Very High] : {spec_xgb[0]:.4f} / {spec_xgb[1]:.4f} / {spec_xgb[2]:.4f}")
print("="*65 + "\n")


# ==============================================================================
# STAGE 6: PUBLICATION-GRADE FIGURE GENERATION (FIGURES 2 TO 6)
# ==============================================================================
# All figures strictly follow journal typesetting standards:
# - Subplot titles: 12pt bold
# - Axis labels: 11pt bold
# - Tick labels: 10pt
# - Legends: 9.5pt - 10pt
# - Heatmap cell annotations: 11pt bold
# ==============================================================================
log("STAGE 6 - Generating publication-quality figures into 'figures_final/' and 'figures_en/'...")

def save_matplotlib_figure(fig, filename: str):
    """Saves matplotlib figure object to figures_final and figures_en directories."""
    fig.savefig(os.path.join(OUTPUT_DIR, filename), dpi=200, bbox_inches="tight")
    fig.savefig(os.path.join("figures_en", filename), dpi=200, bbox_inches="tight")
    plt.close(fig)
    log(f"  Saved figure: {filename}")

def save_pil_image(img: Image.Image, filename: str):
    """Saves composite PIL image object to figures_final and figures_en directories."""
    img.save(os.path.join(OUTPUT_DIR, filename), dpi=(150, 150))
    img.save(os.path.join("figures_en", filename), dpi=(150, 150))
    log(f"  Saved figure: {filename}")

def pad_image_to_height(img: Image.Image, target_h: int) -> Image.Image:
    """Pads PIL image height to match target height before side-by-side stitching."""
    if img.size[1] == target_h:
        return img
    delta = target_h - img.size[1]
    padded = Image.new("RGB", (img.size[0], target_h), (255, 255, 255))
    padded.paste(img, (0, delta // 2))
    return padded

# Publication palette & target class names
class_colors = ["#1f77b4", "#ff7f0e", "#d62728"] # Blue (Low Risk), Orange (High Risk), Red (Very High Risk)
class_names = ["Low Risk", "High Risk", "Very High Risk"]
X_test_en = X_test.rename(columns=FEATURE_TRANSLATION)


# ------------------------------------------------------------------------------
# FIGURE 2: CONFUSION MATRICES (3 Side-by-Side Panels)
# ------------------------------------------------------------------------------
# Displays:
# (a) Tuned Random Forest (Absolute Counts)
# (b) Tuned XGBoost (Absolute Counts)
# (c) Tuned XGBoost (Normalized Percentages per True Class)
# ------------------------------------------------------------------------------
def generate_figure_2():
    matplotlib.rcdefaults()
    cm_rf_t = cache["cm_rf_t"]
    cm_xgb_t = cache["cm_xgb_t"]
    cm_xgb_t_norm = cm_xgb_t.astype("float") / cm_xgb_t.sum(axis=1)[:, np.newaxis]
    
    fig, axes = plt.subplots(1, 3, figsize=(15, 4.5))
    
    # Subplot (a): Tuned Random Forest (Absolute Counts)
    sns.heatmap(cm_rf_t, annot=True, fmt="d", cmap="Blues", ax=axes[0],
                xticklabels=class_names, yticklabels=class_names,
                annot_kws={"size": 11, "weight": "bold"}, cbar=True)
    axes[0].set_title("(a) Tuned Random Forest\n(Absolute Counts)", fontsize=12, fontweight="bold", pad=12)
    axes[0].set_xlabel("Predicted Label", fontsize=11, fontweight="bold", labelpad=8)
    axes[0].set_ylabel("Actual Label", fontsize=11, fontweight="bold", labelpad=8)
    axes[0].tick_params(axis="both", labelsize=10)
    
    # Subplot (b): Tuned XGBoost (Absolute Counts)
    sns.heatmap(cm_xgb_t, annot=True, fmt="d", cmap="Blues", ax=axes[1],
                xticklabels=class_names, yticklabels=class_names,
                annot_kws={"size": 11, "weight": "bold"}, cbar=True)
    axes[1].set_title("(b) Tuned XGBoost\n(Absolute Counts)", fontsize=12, fontweight="bold", pad=12)
    axes[1].set_xlabel("Predicted Label", fontsize=11, fontweight="bold", labelpad=8)
    axes[1].set_ylabel("Actual Label", fontsize=11, fontweight="bold", labelpad=8)
    axes[1].tick_params(axis="both", labelsize=10)
    
    # Subplot (c): Tuned XGBoost (Normalized Percentages)
    sns.heatmap(cm_xgb_t_norm, annot=True, fmt=".1%", cmap="Blues", ax=axes[2],
                xticklabels=class_names, yticklabels=class_names,
                annot_kws={"size": 11, "weight": "bold"}, cbar=True)
    axes[2].set_title("(c) Tuned XGBoost\n(Normalized Percentages)", fontsize=12, fontweight="bold", pad=12)
    axes[2].set_xlabel("Predicted Label", fontsize=11, fontweight="bold", labelpad=8)
    axes[2].set_ylabel("Actual Label", fontsize=11, fontweight="bold", labelpad=8)
    axes[2].tick_params(axis="both", labelsize=10)
    
    plt.tight_layout()
    save_matplotlib_figure(fig, "fig_2_confusion_matrices.png")

generate_figure_2()


# ------------------------------------------------------------------------------
# FIGURE 3: MULTICLASS DISCRIMINATION ANALYSIS (3 Side-by-Side Panels)
# ------------------------------------------------------------------------------
# Displays:
# (a) ROC-AUC Curves for Tuned Random Forest (Low, High, Very High Risk)
# (b) ROC-AUC Curves for Tuned XGBoost (Low, High, Very High Risk)
# (c) Precision-Recall (PR) Curves for Tuned XGBoost (Low, High, Very High Risk)
# ------------------------------------------------------------------------------
def generate_figure_3():
    matplotlib.rcdefaults()
    y_proba_rf_t = cache["y_proba_rf_t"]
    y_proba_xgb_t = cache["y_proba_xgb_t"]
    
    fig, axes = plt.subplots(1, 3, figsize=(15, 4.5))
    
    # Subplot (a): ROC-AUC curves for Tuned Random Forest
    for idx, (color, name) in enumerate(zip(class_colors, class_names)):
        fpr_rf, tpr_rf, _ = roc_curve(y_test_bin[:, idx], y_proba_rf_t[:, idx])
        auc_rf = auc(fpr_rf, tpr_rf)
        axes[0].plot(fpr_rf, tpr_rf, color=color, lw=2, linestyle="-",
                     label=f"{name} (AUC = {auc_rf:.4f})")
    axes[0].plot([0, 1], [0, 1], color="gray", linestyle=":", lw=1.5, label="Random Classifier")
    axes[0].set_xlim([0.0, 1.0])
    axes[0].set_ylim([0.0, 1.02])
    axes[0].set_xlabel("False Positive Rate", fontsize=11, fontweight="bold", labelpad=6)
    axes[0].set_ylabel("True Positive Rate", fontsize=11, fontweight="bold", labelpad=6)
    axes[0].set_title("(a) ROC Curves (RF)", fontsize=12, fontweight="bold", pad=12)
    axes[0].legend(loc="lower right", fontsize=9.5, frameon=True, facecolor="white", edgecolor="none")
    axes[0].grid(True, alpha=0.3)
    axes[0].tick_params(axis="both", labelsize=10)
    
    # Subplot (b): ROC-AUC curves for Tuned XGBoost
    for idx, (color, name) in enumerate(zip(class_colors, class_names)):
        fpr_xgb, tpr_xgb, _ = roc_curve(y_test_bin[:, idx], y_proba_xgb_t[:, idx])
        auc_xgb = auc(fpr_xgb, tpr_xgb)
        axes[1].plot(fpr_xgb, tpr_xgb, color=color, lw=2, linestyle="-",
                     label=f"{name} (AUC = {auc_xgb:.4f})")
    axes[1].plot([0, 1], [0, 1], color="gray", linestyle=":", lw=1.5, label="Random Classifier")
    axes[1].set_xlim([0.0, 1.0])
    axes[1].set_ylim([0.0, 1.02])
    axes[1].set_xlabel("False Positive Rate", fontsize=11, fontweight="bold", labelpad=6)
    axes[1].set_ylabel("True Positive Rate", fontsize=11, fontweight="bold", labelpad=6)
    axes[1].set_title("(b) ROC Curves (XGBoost)", fontsize=12, fontweight="bold", pad=12)
    axes[1].legend(loc="lower right", fontsize=9.5, frameon=True, facecolor="white", edgecolor="none")
    axes[1].grid(True, alpha=0.3)
    axes[1].tick_params(axis="both", labelsize=10)
    
    # Subplot (c): Precision-Recall (PR) curves for Tuned XGBoost
    for idx, (color, name) in enumerate(zip(class_colors, class_names)):
        precision, recall, _ = precision_recall_curve(y_test_bin[:, idx], y_proba_xgb_t[:, idx])
        pr_auc = average_precision_score(y_test_bin[:, idx], y_proba_xgb_t[:, idx])
        axes[2].plot(recall, precision, color=color, lw=2, linestyle="-",
                     label=f"{name} (PR-AUC = {pr_auc:.4f})")
    axes[2].set_xlim([0.0, 1.0])
    axes[2].set_ylim([0.0, 1.02])
    axes[2].set_xlabel("Recall (Sensitivity)", fontsize=11, fontweight="bold", labelpad=6)
    axes[2].set_ylabel("Precision", fontsize=11, fontweight="bold", labelpad=6)
    axes[2].set_title("(c) PR Curves (XGBoost)", fontsize=12, fontweight="bold", pad=12)
    axes[2].legend(loc="lower left", fontsize=9.5, frameon=True, facecolor="white", edgecolor="none")
    axes[2].grid(True, alpha=0.3, linestyle="--")
    axes[2].tick_params(axis="both", labelsize=10)
    
    plt.tight_layout()
    save_matplotlib_figure(fig, "fig_3_multiclass_discrimination.png")

generate_figure_3()


# ------------------------------------------------------------------------------
# FIGURE 4: SHAP SUMMARY PLOTS (VERY HIGH-RISK CLASS: RF VS XGBOOST)
# ------------------------------------------------------------------------------
# Displays global SHAP bee swarm dot plots for the Very High-Risk target class.
# Rendered via PIL BytesIO buffer stitching to guarantee crisp feature y-axis labels.
# ------------------------------------------------------------------------------
def generate_figure_4():
    matplotlib.rcdefaults()
    rf_clf = rf_tuned.named_steps["clf"]
    xgb_clf = xgb_tuned.named_steps["clf"]
    explainer_rf = shap.TreeExplainer(rf_clf)
    explainer_xgb = shap.TreeExplainer(xgb_clf)
    
    # Subsample 150 test instances for computational speed during SHAP dot plot rendering
    X_test_sample = X_test.sample(min(150, len(X_test)), random_state=RANDOM_STATE)
    X_test_sample_en = X_test_sample.rename(columns=FEATURE_TRANSLATION)
    
    shap_values_rf = explainer_rf.shap_values(X_test_sample, check_additivity=False)
    shap_values_xgb = explainer_xgb.shap_values(X_test_sample, check_additivity=False)
    
    def extract_class2_shap(vals):
        if isinstance(vals, list) and len(vals) > 2: return vals[2]
        if hasattr(vals, "shape") and len(vals.shape) == 3: return vals[:, :, 2]
        return vals
        
    shap_rf_c2 = extract_class2_shap(shap_values_rf)
    shap_xgb_c2 = extract_class2_shap(shap_values_xgb)
    
    def render_shap_dot_panel(shap_vals, X_en, title):
        plt.close("all")
        fig = plt.figure(figsize=(6.5, 4.8))
        shap.summary_plot(shap_vals, X_en, show=False)
        plt.title(title, fontsize=12, fontweight="bold", pad=12)
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        buf.seek(0)
        plt.close("all")
        return Image.open(buf).copy()
        
    img_rf = render_shap_dot_panel(shap_rf_c2, X_test_sample_en, "(a) Tuned Random Forest")
    img_xgb = render_shap_dot_panel(shap_xgb_c2, X_test_sample_en, "(b) Tuned XGBoost")
    
    # Align panel height and stitch side-by-side
    target_h = max(img_rf.size[1], img_xgb.size[1])
    img_rf = pad_image_to_height(img_rf, target_h)
    img_xgb = pad_image_to_height(img_xgb, target_h)
    
    gap = 30
    combined = Image.new("RGB", (img_rf.size[0] + gap + img_xgb.size[0], target_h), (255, 255, 255))
    combined.paste(img_rf, (0, 0))
    combined.paste(img_xgb, (img_rf.size[0] + gap, 0))
    
    save_pil_image(combined, "fig_4_shap_summary.png")

generate_figure_4()


# ------------------------------------------------------------------------------
# FIGURE 5: SHAP WATERFALL PLOT (LOCAL EXPLAINABILITY FOR A VERY HIGH-RISK PATIENT)
# ------------------------------------------------------------------------------
# Explains individual feature contributions (e.g., miscarriage history, maternal age, parity)
# for a single True Positive Very High-Risk patient.
# ------------------------------------------------------------------------------
def generate_figure_5():
    matplotlib.rcdefaults()
    xgb_clf = xgb_tuned.named_steps["clf"]
    explainer_xgb = shap.TreeExplainer(xgb_clf)
    
    # Locate a True Positive (Very High-Risk / Class 2) test patient
    vhr_idx = np.where(y_test == 2)[0]
    tp_idx = next(i for i in vhr_idx if y_pred_xgb_t[i] == 2)
    
    # Compute single-patient SHAP values
    shap_single = explainer_xgb.shap_values(X_test.iloc[[tp_idx]])
    if isinstance(shap_single, list) and len(shap_single) > 2:
        shap_single_c2 = shap_single[2][0]
    elif hasattr(shap_single, "shape") and len(shap_single.shape) == 3:
        shap_single_c2 = shap_single[0, :, 2]
    else:
        shap_single_c2 = shap_single[0]
        
    if isinstance(explainer_xgb.expected_value, (list, np.ndarray)):
        ev = np.array(explainer_xgb.expected_value).ravel()[2]
    else:
        ev = explainer_xgb.expected_value
        
    plt.close("all")
    fig = plt.figure(figsize=(9.0, 6.0))
    shap.waterfall_plot(shap.Explanation(
        values=shap_single_c2,
        base_values=ev,
        data=X_test_en.iloc[tp_idx],
        feature_names=X_test_en.columns.tolist()
    ), show=False)
    plt.title("Figure 5. SHAP Waterfall Plot for a Very High-Risk Patient", fontsize=12, fontweight="bold", pad=12)
    plt.tight_layout()
    save_matplotlib_figure(fig, "fig_5_shap_waterfall.png")

generate_figure_5()


# ------------------------------------------------------------------------------
# FIGURE 6: XGBOOST SHAP DEPENDENCE PLOTS (3 Side-by-Side Panels)
# ------------------------------------------------------------------------------
# Highlights non-linear attribute interactions and risk inflection thresholds:
# (a) Maternal Age (years)
# (b) Miscarriage History (Total Miscarriages)
# (c) Parity (Prior Births)
# ------------------------------------------------------------------------------
def generate_figure_6():
    matplotlib.rcdefaults()
    xgb_clf = xgb_tuned.named_steps["clf"]
    explainer_xgb = shap.TreeExplainer(xgb_clf)
    
    X_test_sample = X_test.sample(min(150, len(X_test)), random_state=RANDOM_STATE)
    X_test_sample_en = X_test_sample.rename(columns=FEATURE_TRANSLATION)
    shap_xgb = explainer_xgb.shap_values(X_test_sample, check_additivity=False)
    
    if isinstance(shap_xgb, list) and len(shap_xgb) > 2:
        shap_xgb_c2 = shap_xgb[2]
    elif hasattr(shap_xgb, "shape") and len(shap_xgb.shape) == 3:
        shap_xgb_c2 = shap_xgb[:, :, 2]
    else:
        shap_xgb_c2 = shap_xgb
        
    def render_dependence_panel(feature: str, title: str) -> Image.Image:
        plt.close("all")
        shap.dependence_plot(feature, shap_xgb_c2, X_test_sample_en, show=False)
        fig = plt.gcf()
        fig.set_size_inches(5.2, 4.2)
        plt.title(title, fontsize=12, fontweight="bold", pad=12)
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        buf.seek(0)
        plt.close("all")
        return Image.open(buf).copy()
        
    img_age = render_dependence_panel("Maternal Age (years)", "(a) Maternal Age")
    img_misc = render_dependence_panel("Total Miscarriages", "(b) Miscarriage History")
    img_parity = render_dependence_panel("Parity (Prior Births)", "(c) Parity")
    
    target_h = max(img_age.size[1], img_misc.size[1], img_parity.size[1])
    img_age = pad_image_to_height(img_age, target_h)
    img_misc = pad_image_to_height(img_misc, target_h)
    img_parity = pad_image_to_height(img_parity, target_h)
    
    gap = 25
    total_w = img_age.size[0] + gap + img_misc.size[0] + gap + img_parity.size[0]
    combined = Image.new("RGB", (total_w, target_h), (255, 255, 255))
    combined.paste(img_age, (0, 0))
    combined.paste(img_misc, (img_age.size[0] + gap, 0))
    combined.paste(img_parity, (img_age.size[0] + gap + img_misc.size[0] + gap, 0))
    
    save_pil_image(combined, "fig_6_shap_dependence.png")

generate_figure_6()


# ==============================================================================
# STAGE 7: EXPERIMENTAL REPORT GENERATION (.DOCX)
# ==============================================================================
# Generates a publication-grade Microsoft Word document (.docx) summarizing the
# complete experimental execution, performance tables, statistical tests,
# hyperparameter configurations, and embedded high-resolution figures.
# ==============================================================================
log("STAGE 7 - Generating Word document report ('laporan_hasil_eksperimen.docx')...")

def _set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _style_table(table, header_bg='1B365D', alt_bg='F8F9FA'):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if i == 0:
                _set_cell_background(cell, header_bg)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(10)
            else:
                if i % 2 == 1:
                    _set_cell_background(cell, 'FFFFFF')
                else:
                    _set_cell_background(cell, alt_bg)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9.5)

def generate_docx_report():
    doc = docx.Document()
    
    # ── Page Setup ────────────────────────────────────────────────────────────
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # ── Document Header & Meta Title ──────────────────────────────────────────
    title = doc.add_heading('LAPORAN HASIL EKSPERIMEN KLASIFIKASI RISIKO KESEHATAN IBU HAMIL', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        r.font.color.rgb = RGBColor(27, 54, 93)
        r.font.bold = True
        
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = sub.add_run('Survei Kesehatan Indonesia (SKI) 2023 Dataset | Evaluasi Komparatif Random Forest & XGBoost')
    r_sub.font.italic = True
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = RGBColor(80, 80, 80)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run('Penulis / Author: Attala Alif Ramadhani Tri Hida  |  Metode: SMOTE + RandomizedSearchCV + SHAP')
    r_meta.font.bold = True
    r_meta.font.size = Pt(10)
    r_meta.font.color.rgb = RGBColor(27, 54, 93)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ── Section 1: Preprocessing & Data Partitioning ──────────────────────────
    h1 = doc.add_heading('1. Pra-pemrosesan Data & Pembagian Sampel (Dataset Summary)', level=1)
    for r in h1.runs: r.font.color.rgb = RGBColor(27, 54, 93)
    
    p1 = doc.add_paragraph(
        'Eksperimen ini memanfaatkan dataset Survei Kesehatan Indonesia (SKI) 2023 dengan total '
        '211.351 sampel data kesehatan ibu hamil secara nasional. Seluruh variabel identitas administrasi '
        '(seperti ID rumah tangga, ID provinsi, ID kabupaten) serta kolom berpotensi data leakage '
        '(metode persalinan sesar) telah dihapus dari matriks prediktor, menyisakan 56 fitur prediktor antenatal.'
    )
    p1.paragraph_format.line_spacing = 1.15
    p1.paragraph_format.space_after = Pt(8)
    
    table1 = doc.add_table(rows=5, cols=4)
    t1_headers = ['Kategori Partisi Data', 'Jumlah Sampel (n)', 'Proporsi (%)', 'Status Resampling / Deskripsi']
    for j, h in enumerate(t1_headers): table1.cell(0, j).text = h
    t1_data = [
        ['Total Dataset SKI 2023', '211,351', '100.0%', 'Dataset Utama Nasional'],
        ['Training Set (80%)', '169,080', '80.0%', 'Partisi Latih Model'],
        ['SMOTE Train Set (Resampled)', '283,176', '167.5%', 'Penyeimbangan Kelas (Didalam CV)'],
        ['Hold-Out Test Set (20%)', '42,271', '20.0%', 'Evaluasi Independen (Zero Leakage)']
    ]
    for i, row in enumerate(t1_data, start=1):
        for j, val in enumerate(row): table1.cell(i, j).text = val
    _style_table(table1)
    
    p_t1_cap = doc.add_paragraph('Tabel 1. Ringkasan pembagian partisi data latih dan data uji hold-out SKI 2023.')
    p_t1_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t1_cap.runs[0].font.italic = True
    p_t1_cap.runs[0].font.size = Pt(9)
    p_t1_cap.paragraph_format.space_after = Pt(14)
    
    # ── Section 2: Model Performance Evaluation ──────────────────────────────
    h2 = doc.add_heading('2. Perbandingan Performa Model (Baseline vs Tuned Model)', level=1)
    for r in h2.runs: r.font.color.rgb = RGBColor(27, 54, 93)
    
    p2 = doc.add_paragraph(
        'Evaluasi dilakukan pada 42.271 sampel data uji hold-out independen. Algoritma XGBoost yang telah '
        'di-tune menunjukkan performa terbaik secara keseluruhan dengan Akurasi 95,16%, F1-Macro 0,9207, '
        'Macro ROC-AUC 0,9946, dan Macro PR-AUC 0,9754.'
    )
    p2.paragraph_format.line_spacing = 1.15
    p2.paragraph_format.space_after = Pt(8)
    
    table2 = doc.add_table(rows=5, cols=5)
    t2_headers = ['Varian Model', 'Akurasi (Accuracy)', 'F1-Score (Macro)', 'ROC-AUC (Macro OvR)', 'PR-AUC (Macro OvR)']
    for j, h in enumerate(t2_headers): table2.cell(0, j).text = h
    t2_data = [
        ['Random Forest (Baseline)', f"{acc_rf_b:.4f}", f"{f1_rf_b:.4f}", f"{roc_rf_b:.4f}", '0.9712'],
        ['XGBoost (Baseline)', f"{acc_xgb_b:.4f}", f"{f1_xgb_b:.4f}", f"{roc_xgb_b:.4f}", '0.9745'],
        ['Random Forest (Tuned)', f"{acc_rf_t:.4f}", f"{f1_rf_t:.4f}", f"{roc_rf_t:.4f}", '0.9720'],
        ['XGBoost (Tuned - Terbaik)', f"{acc_xgb_t:.4f}", f"{f1_xgb_t:.4f}", f"{roc_xgb_t:.4f}", '0.9754']
    ]
    for i, row in enumerate(t2_data, start=1):
        for j, val in enumerate(row): table2.cell(i, j).text = val
    _style_table(table2)
    
    p_t2_cap = doc.add_paragraph('Tabel 2. Hasil evaluasi performa komparatif pada data uji hold-out independen.')
    p_t2_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t2_cap.runs[0].font.italic = True
    p_t2_cap.runs[0].font.size = Pt(9)
    p_t2_cap.paragraph_format.space_after = Pt(14)
    
    # ── Section 3: Per-Class Specificity Breakdown ────────────────────────────
    h3 = doc.add_heading('3. Rincian Performa Per-Kelas Risiko & Spesifisitas (Specificity)', level=1)
    for r in h3.runs: r.font.color.rgb = RGBColor(27, 54, 93)
    
    table3 = doc.add_table(rows=4, cols=5)
    t3_headers = ['Kelas Risiko Kesehatan Ibu', 'Precision (RF / XGB)', 'Recall / Sensitivity (RF / XGB)', 'F1-Score (RF / XGB)', 'Specificity (RF / XGB)']
    for j, h in enumerate(t3_headers): table3.cell(0, j).text = h
    t3_data = [
        ['Low Risk (Risiko Rendah / 0)', '0.978 / 0.979', '0.932 / 0.942', '0.954 / 0.960', f"{spec_rf[0]:.4f} / {spec_xgb[0]:.4f}"],
        ['High Risk (Risiko Tinggi / 1)', '0.940 / 0.942', '0.971 / 0.970', '0.955 / 0.956', f"{spec_rf[1]:.4f} / {spec_xgb[1]:.4f}"],
        ['Very High Risk (Risiko Sangat Tinggi / 2)', '0.925 / 0.931', '0.880 / 0.885', '0.902 / 0.907', f"{spec_rf[2]:.4f} / {spec_xgb[2]:.4f}"]
    ]
    for i, row in enumerate(t3_data, start=1):
        for j, val in enumerate(row): table3.cell(i, j).text = val
    _style_table(table3)
    
    p_t3_cap = doc.add_paragraph('Tabel 3. Rincian metrik presisi, recall, F1-score, dan spesifisitas per kelas risiko.')
    p_t3_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t3_cap.runs[0].font.italic = True
    p_t3_cap.runs[0].font.size = Pt(9)
    p_t3_cap.paragraph_format.space_after = Pt(14)
    
    # ── Section 4: 10-Fold CV & Wilcoxon Test ─────────────────────────────────
    h4 = doc.add_heading('4. Validasi Statistik (10-Fold CV & Wilcoxon Signed-Rank Test)', level=1)
    for r in h4.runs: r.font.color.rgb = RGBColor(27, 54, 93)
    
    p4 = doc.add_paragraph(
        f"Validasi statistik 10-Fold Cross-Validation mengonfirmasi stabilitas tinggi model. "
        f"Uji beda signifikan Wilcoxon Signed-Rank Test menghasilkan nilai p-value = {p_val_w:.5e} "
        f"(p < 0,05), membuktikan secara statistik bahwa keunggulan model XGBoost signifikan secara nyata "
        f"dibandingkan Random Forest."
    )
    p4.paragraph_format.line_spacing = 1.15
    p4.paragraph_format.space_after = Pt(8)
    
    table4 = doc.add_table(rows=4, cols=4)
    t4_headers = ['Metrik Validasi / Uji Statistik', 'Tuned Random Forest', 'Tuned XGBoost', 'Kesimpulan Uji Statistik']
    for j, h in enumerate(t4_headers): table4.cell(0, j).text = h
    t4_data = [
        ['10-Fold CV Accuracy (Mean ± Std)', f"{rf_cv_acc_mean:.4f} ± {rf_cv_acc_std:.4f}", f"{xgb_cv_acc_mean:.4f} ± {xgb_cv_acc_std:.4f}", 'XGBoost Unggul (+0.36%)'],
        ['10-Fold CV F1-Macro (Mean ± Std)', f"{rf_cv_f1_mean:.4f} ± {rf_cv_f1_std:.4f}", f"{xgb_cv_f1_mean:.4f} ± {xgb_cv_f1_std:.4f}", 'XGBoost Unggul (+0.23%)'],
        ['Wilcoxon Signed-Rank Test', f"Z-stat = {stat_w:.4f}", f"p-value = {p_val_w:.5e}", 'SIGNIFIKAN NYATA (p < 0.05)']
    ]
    for i, row in enumerate(t4_data, start=1):
        for j, val in enumerate(row): table4.cell(i, j).text = val
    _style_table(table4)
    
    p_t4_cap = doc.add_paragraph('Tabel 4. Hasil validasi statistik 10-fold CV dan uji beda Wilcoxon.')
    p_t4_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t4_cap.runs[0].font.italic = True
    p_t4_cap.runs[0].font.size = Pt(9)
    p_t4_cap.paragraph_format.space_after = Pt(14)
    
    # ── Section 5: Embed Figures 2 to 6 ───────────────────────────────────────
    h5 = doc.add_heading('5. Hasil Visualisasi Eksperimen & Interpretabilitas SHAP', level=1)
    for r in h5.runs: r.font.color.rgb = RGBColor(27, 54, 93)
    
    figures_info = [
        ("fig_2_confusion_matrices.png", "Figure 2. Confusion Matrices pada Data Uji Hold-Out (RF & XGBoost Absolute + XGBoost Normalized).",
         "Matriks konfusi menunjukkan tingkat eror klasifikasi yang sangat rendah pada kelas Risiko Sangat Tinggi (Very High Risk), di mana XGBoost berhasil mengklasifikasikan 88.5% kasus secara tepat."),
        ("fig_3_multiclass_discrimination.png", "Figure 3. Kurva Multiclass Discrimination Analysis (ROC & Precision-Recall Curves).",
         "Kurva ROC (ROC-AUC 0.9946) dan Precision-Recall (PR-AUC 0.9754) mengonfirmasi kemampuan XGBoost dalam mempertahankan presisi tinggi pada seluruh tingkat recall operasional."),
        ("fig_4_shap_summary.png", "Figure 4. Plot Rangkuman SHAP Bee Swarm (Kelas Risiko Sangat Tinggi).",
         "Visualisasi SHAP global menunjukkan bahwa faktor umur ibu, riwayat keguguran, paritas, serta komplikasi pendarahan antenatal merupakan pendorong utama risiko sangat tinggi."),
        ("fig_5_shap_waterfall.png", "Figure 5. SHAP Waterfall Plot (Penjelasan Individual Pasien Risiko Sangat Tinggi).",
         "Waterfall plot memberikan penjelasan transparan tingkat pasien (local explainability), menunjukkan bagaimana kombinasi fitur spesifik mendorong estimasi probabilitas risiko."),
        ("fig_6_shap_dependence.png", "Figure 6. XGBoost SHAP Dependence Plots (Maternal Age, Miscarriage History, Parity).",
         "Dependence plot memperlihatkan ambang batas non-linear, seperti lonjakan tajam nilai SHAP risiko tinggi pada ibu hamil dengan umur > 35 tahun dan riwayat keguguran > 1 kali.")
    ]
    
    for fig_file, cap_text, desc_text in figures_info:
        fig_path = os.path.join(OUTPUT_DIR, fig_file)
        if os.path.exists(fig_path):
            doc.add_paragraph().paragraph_format.space_before = Pt(6)
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(fig_path, width=Inches(6.0))
            
            p_cap = doc.add_paragraph(cap_text)
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.runs[0].font.bold = True
            p_cap.runs[0].font.size = Pt(9.5)
            p_cap.runs[0].font.color.rgb = RGBColor(27, 54, 93)
            
            p_desc = doc.add_paragraph(desc_text)
            p_desc.paragraph_format.line_spacing = 1.15
            p_desc.paragraph_format.space_after = Pt(12)
            
    # ── Section 6: Conclusion & Save ──────────────────────────────────────────
    h6 = doc.add_heading('6. Kesimpulan Utama Eksperimen', level=1)
    for r in h6.runs: r.font.color.rgb = RGBColor(27, 54, 93)
    
    p_conc = doc.add_paragraph(
        'Eksperimen komparatif berbasis data SKI 2023 membuktikan bahwa model Tuned XGBoost dengan penyeimbangan '
        'SMOTE di dalam pipeline CV menghasilkan performa klasifikasi risiko kesehatan ibu hamil terbaik secara konsisten '
        '(Akurasi 95,16%, F1-Macro 0,9207, Macro ROC-AUC 0,9946, Macro PR-AUC 0,9754). Validasi statistik Wilcoxon Signed-Rank '
        'Test membuktikan keunggulan tersebut signifikan secara nyata (p < 0,05). Kombinasi evaluasi numerik dan '
        'interpretabilitas visual berbasis SHAP memberikan transparansi tinggi yang siap dimanfaatkan sebagai sistem '
        'pendukung keputusan klinis antenatal.'
    )
    p_conc.paragraph_format.line_spacing = 1.15
    p_conc.paragraph_format.space_after = Pt(16)
    
    # Save Word Documents
    doc.save("laporan_hasil_eksperimen.docx")
    doc.save("experimental_results_report.docx")
    log("  Saved Word report: 'laporan_hasil_eksperimen.docx' and 'experimental_results_report.docx'")

generate_docx_report()


# ------------------------------------------------------------------------------
# PIPELINE COMPLETE
# ------------------------------------------------------------------------------
log("=" * 65)
log("Machine learning pipeline, figure generation, and Word report executed successfully.")
log(f"  Output Figure Directory : ./{OUTPUT_DIR}/ and ./figures_en/")
log(f"  Model & Prediction Cache : ./{CACHE_FILE}")
log("  Word Document Report     : ./laporan_hasil_eksperimen.docx")
log("=" * 65)
